from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"C:\Users\lucas\Desktop\Bryza\docs\Playbook_Programa_de_Embaixadores_Bryza.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "005675"
BLUE = "2A6F8F"
LIGHT_BLUE = "EAF5FA"
PALE_BLUE = "F4F9FC"
GOLD = "C88A20"
PALE_GOLD = "FFF6E2"
GREEN = "287A59"
PALE_GREEN = "EAF6F0"
RED = "A63D40"
PALE_RED = "FCEEEF"
INK = "24323B"
MUTED = "5F6F78"
LINE = "C9DCE5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[i]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    paragraph.add_run("Página ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_numbering(doc, num_id, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_id = 900 + num_id
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])
    ppr.append(num_pr)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_para(doc, text="", bold_prefix=None, style=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = add_para(doc, after=4)
        apply_num(p, 41)
        if isinstance(item, tuple):
            p.add_run(item[0]).bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = add_para(doc, after=5)
        apply_num(p, 42)
        if isinstance(item, tuple):
            p.add_run(item[0]).bold = True
            p.add_run(item[1])
        else:
            p.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_callout(doc, label, text, kind="info"):
    palette = {
        "info": (LIGHT_BLUE, NAVY),
        "success": (PALE_GREEN, GREEN),
        "warning": (PALE_GOLD, GOLD),
        "danger": (PALE_RED, RED),
    }
    fill, accent = palette[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    borders.append(left)
    ppr.append(borders)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(str(value))
    set_table_geometry(table, widths, 120)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.8)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for level, size, before, after, color in [
    (1, 16, 18, 10, NAVY),
    (2, 13, 14, 7, BLUE),
    (3, 11.5, 10, 5, NAVY),
]:
    style = styles[f"Heading {level}"]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title_style = styles["Title"]
title_style.font.name = "Calibri"
title_style.font.size = Pt(31)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(NAVY)
title_style.paragraph_format.space_after = Pt(10)

subtitle = styles.add_style("Bryza Subtitle", WD_STYLE_TYPE.PARAGRAPH)
subtitle.font.name = "Calibri"
subtitle.font.size = Pt(14)
subtitle.font.color.rgb = RGBColor.from_string(BLUE)
subtitle.paragraph_format.space_after = Pt(12)

add_numbering(doc, 41, bullet=True)
add_numbering(doc, 42, bullet=False)

# Header/footer
for sec in doc.sections:
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("BRYZA  |  PROGRAMA DE EMBAIXADORES")
    hr.font.name = "Calibri"
    hr.font.size = Pt(8)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor.from_string(MUTED)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    add_page_field(fp)
    for run in fp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(40)
k = doc.add_paragraph()
k.alignment = WD_ALIGN_PARAGRAPH.CENTER
k.paragraph_format.space_after = Pt(18)
kr = k.add_run("PLAYBOOK OPERACIONAL")
kr.bold = True
kr.font.size = Pt(11)
kr.font.color.rgb = RGBColor.from_string(GOLD)

t = doc.add_paragraph("Programa de Embaixadores Bryza", style="Title")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = doc.add_paragraph("Como funciona, como explicar, como ativar e como operar o sistema de ponta a ponta", style="Bryza Subtitle")
s.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(12)
add_callout(doc, "Versão verificada", "21 de julho de 2026 • baseada na configuração vigente e na implementação do sistema Bryza.", "info")

doc.add_paragraph().paragraph_format.space_after = Pt(30)
q = doc.add_paragraph()
q.alignment = WD_ALIGN_PARAGRAPH.CENTER
q.paragraph_format.space_after = Pt(12)
qr = q.add_run("A ideia em uma frase")
qr.bold = True
qr.font.size = Pt(12)
qr.font.color.rgb = RGBColor.from_string(NAVY)
qq = doc.add_paragraph("Cada pessoa recebe um código de indicação, forma uma rede ao apresentar novos embaixadores e pode ganhar comissões sobre vendas elegíveis em até três níveis — sempre conforme o plano congelado no pedido e os status operacionais.")
qq.alignment = WD_ALIGN_PARAGRAPH.CENTER
qq.paragraph_format.line_spacing = 1.3

doc.add_page_break()

add_heading(doc, "1. Resumo executivo", 1)
add_callout(doc, "Regra vigente", "Plano Embaixador Multinível 4–2–1: 4% no nível 1, 2% no nível 2 e 1% no nível 3, calculados sobre o valor final comissionável da venda.", "success")
add_table(doc, ["Parâmetro", "Configuração atual"], [
    ("Status do programa", "Ativo"),
    ("Plano padrão", "Embaixador Multinível 4–2–1"),
    ("Base de cálculo", "Valor final da venda (valor comissionável congelado no pedido)"),
    ("Janela da indicação", "30 dias a partir da visita válida"),
    ("Pagamento mínimo", "R$ 50,00 em comissões liberadas"),
    ("Frequência operacional", "Mensal"),
    ("Meio de repasse", "Pix, registrado manualmente pelo administrador"),
    ("Ativação mensal configurada", "R$ 79,00 em compras pessoais + 10 dias de tolerância"),
    ("Bônus da 1ª compra", "R$ 7,00 quando a primeira compra elegível for de pelo menos R$ 79,00"),
], [2500, 6740])

add_heading(doc, "O que o embaixador precisa entender primeiro", 2)
add_bullets(doc, [
    ("Código e link próprios. ", "O sistema gera um identificador no formato bryzaNN, usado também no link e no login."),
    ("Venda direta. ", "Quando alguém compra pelo link, o embaixador indicador ocupa o nível 1 daquela venda."),
    ("Rede. ", "Quem entra a partir desse comprador fica ligado ao indicador como patrocinado; a cadeia pode remunerar os níveis 2 e 3."),
    ("Comissão não é pagamento imediato. ", "Ela nasce vinculada ao pedido, passa por estados de validação e só pode ser paga quando estiver liberada."),
    ("Regra histórica é preservada. ", "Plano, percentuais e base são congelados no pedido; mudanças futuras não reescrevem o passado."),
])

add_heading(doc, "A explicação de 30 segundos", 2)
add_callout(doc, "Script", "“Você recebe um link Bryza. Quando uma pessoa compra por ele, a venda fica registrada na sua indicação. No plano atual, você ganha 4% nas vendas diretas, 2% nas vendas dos embaixadores que você indicou e 1% no terceiro nível. A comissão só é liberada depois da entrega e da confirmação do pagamento. Quando o saldo liberado chega a R$ 50, a Bryza faz o repasse por Pix no ciclo mensal.”", "info")

doc.add_page_break()

add_heading(doc, "2. O sistema de ponta a ponta", 1)
add_steps(doc, [
    ("Compartilhamento. ", "O embaixador envia seu link ou QR Code."),
    ("Visita rastreada. ", "O acesso registra a visita e grava um cookie assinado de atribuição."),
    ("Compra/agendamento. ", "O cliente informa seus dados, escolhe produtos e agenda a entrega; a indicação válida é congelada."),
    ("Novo código. ", "O comprador recebe automaticamente um código próprio e passa a existir como candidato a embaixador, ainda pendente de acesso."),
    ("Conversão em pedido. ", "A operação transforma o agendamento em pedido sem perder a origem, o plano e a cadeia da indicação."),
    ("Comissões geradas. ", "O sistema calcula os beneficiários ativos dos níveis 1, 2 e 3."),
    ("Entrega e pagamento confirmados. ", "As comissões deixam de aguardar e tornam-se liberadas."),
    ("Repasse. ", "O administrador seleciona as comissões liberadas, registra a referência do Pix e conclui o pagamento."),
])

add_heading(doc, "Atribuição: quem fica com a indicação?", 2)
add_bullets(doc, [
    "Somente links de embaixadores com status ativo são aceitos.",
    "A visita precisa existir e estar dentro da janela configurada de 30 dias.",
    "O cookie de atribuição é assinado; alterações manuais no navegador não devem ser aceitas.",
    "A primeira atribuição confirmada do cliente prevalece e fica bloqueada para compras seguintes.",
    "Um novo clique pode autorizar o checkout, mas não troca automaticamente o primeiro indicador já bloqueado.",
    "A administração pode reatribuir antes de estados terminais, com motivo e trilha de auditoria.",
])
add_callout(doc, "Importante", "A atribuição é por cliente, não apenas por uma compra isolada. Depois de bloqueada, ela tende a acompanhar o relacionamento do cliente com a Bryza.", "warning")

add_heading(doc, "O que fica congelado no pedido", 2)
add_bullets(doc, [
    "Embaixador atribuído e código de indicação.",
    "Plano de comissão vigente e percentuais por nível.",
    "Base comissionável e valor da comissão direta calculada.",
    "Configuração de bônus da primeira compra aplicável naquele momento.",
])

add_heading(doc, "3. Rede multinível: 4–2–1 sem mistério", 1)
add_table(doc, ["Nível", "Quem recebe", "Percentual atual"], [
    ("1", "O embaixador cujo link originou/possui a atribuição do cliente", "4%"),
    ("2", "O patrocinador do embaixador do nível 1", "2%"),
    ("3", "O patrocinador do nível 2", "1%"),
], [1100, 6500, 1640])
add_para(doc, "A cadeia sobe a partir do embaixador diretamente atribuído ao pedido. O sistema admite planos configuráveis de até dez níveis, mas o plano vigente possui três níveis ativos.")

add_heading(doc, "Exemplo simples", 2)
add_table(doc, ["Venda", "Nível 1 (4%)", "Nível 2 (2%)", "Nível 3 (1%)", "Total rede"], [
    ("R$ 79,00", "R$ 3,16", "R$ 1,58", "R$ 0,79", "R$ 5,53"),
    ("R$ 100,00", "R$ 4,00", "R$ 2,00", "R$ 1,00", "R$ 7,00"),
    ("R$ 500,00", "R$ 20,00", "R$ 10,00", "R$ 5,00", "R$ 35,00"),
], [1700, 1900, 1900, 1900, 1840])
add_callout(doc, "Fórmula", "Comissão = base comissionável do pedido × percentual do nível. O resultado é arredondado para duas casas decimais.", "info")

add_heading(doc, "Regras que evitam interpretações erradas", 2)
add_bullets(doc, [
    ("Sem compressão. ", "Se uma pessoa da cadeia estiver inativa, ela não recebe; o percentual não é transferido para outro nível."),
    ("Sem autoindicação. ", "O sistema bloqueia coincidências por vínculo do cliente, CPF ou telefone."),
    ("Somente beneficiário ativo. ", "Cada nó da cadeia precisa estar com status ativo no momento da geração da comissão."),
    ("Sem duplicidade. ", "Uma mesma venda, beneficiário, nível e tipo de comissão não podem gerar o mesmo registro duas vezes."),
    ("Mudança de plano não retroage. ", "Novas regras valem para novos pedidos; o histórico mantém seus snapshots."),
])

doc.add_page_break()

add_heading(doc, "4. Bônus da primeira compra", 1)
add_callout(doc, "Regra atual", "O patrocinador direto recebe um bônus fixo de R$ 7,00 quando a primeira compra elegível do novo indicado for de pelo menos R$ 79,00.", "success")
add_heading(doc, "Todas as condições precisam ser verdadeiras", 2)
add_bullets(doc, [
    "O bônus estava ativo quando o pedido foi criado.",
    "O pedido totaliza pelo menos R$ 79,00.",
    "O pedido chega a entregue ou finalizado e o pagamento é confirmado.",
    "O cliente possui seu próprio cadastro de candidato a embaixador.",
    "Esse candidato está ligado como filho direto do embaixador indicador.",
    "Existe a comissão percentual de nível 1 da mesma venda.",
    "O cliente ainda não recebeu bônus de primeira compra em outro pedido elegível.",
    "O cliente/candidato foi cadastrado depois da ativação da regra de bônus.",
])
add_para(doc, "O bônus entra na mesma carteira de comissões, mas aparece identificado como “Bônus 1ª compra”. Ele é liberado diretamente quando a venda cumpre entrega e confirmação financeira.")

add_heading(doc, "Exemplo combinado em uma compra de R$ 79,00", 2)
add_table(doc, ["Beneficiário", "Componente", "Valor"], [
    ("Patrocinador direto", "4% sobre R$ 79,00", "R$ 3,16"),
    ("Patrocinador direto", "Bônus fixo da 1ª compra", "R$ 7,00"),
    ("Nível 2", "2% sobre R$ 79,00", "R$ 1,58"),
    ("Nível 3", "1% sobre R$ 79,00", "R$ 0,79"),
    ("Total distribuído", "Comissão de rede + bônus", "R$ 12,53"),
], [2800, 4440, 2000])
add_callout(doc, "Não confundir", "O bônus de R$ 7 não substitui os 4% do nível 1; quando todas as condições são atendidas, os dois valores coexistem.", "warning")

add_heading(doc, "5. Ativação: três conceitos diferentes", 1)
add_heading(doc, "5.1 Ativação cadastral e acesso ao painel", 2)
add_table(doc, ["Status", "O que significa"], [
    ("Pendente", "Cadastro existe, mas o login não é permitido."),
    ("Ativo", "Perfil habilitado; pode entrar no painel e pode receber novas comissões se os demais critérios forem atendidos."),
    ("Inativo", "Acesso desabilitado por decisão operacional; não recebe novas comissões durante a inatividade."),
    ("Bloqueado", "Acesso impedido por controle/segurança; não recebe novas comissões."),
], [1700, 7540])
add_para(doc, "No cadastro administrativo completo, a Bryza cria o usuário, gera o código bryzaNN e define a senha inicial como o telefone com DDD, apenas números. No primeiro acesso, o embaixador informa o CPF e cria uma nova senha com no mínimo oito caracteres.")
add_callout(doc, "Fluxo recomendado", "Cadastrar ou completar os dados → conferir plano e patrocinador → validar chave Pix → provisionar o usuário → marcar como ativo → entregar código, link e instruções do primeiro acesso.", "info")

add_heading(doc, "5.2 Cadastro automático pela página de vendas", 2)
add_para(doc, "Quem compra pelo link recebe automaticamente um código e entra na rede como filho do indicador. Esse registro nasce como candidato pendente, com a observação de que o acesso ainda não foi ativado.")
add_callout(doc, "Comportamento atual do sistema", "O checkout cria o candidato e o código, mas não cria uma conta de autenticação. Alterar apenas o status para “ativo” não provisiona usuário nem senha. Portanto, antes de prometer acesso, a operação precisa concluir o provisionamento administrativo.", "danger")

add_heading(doc, "5.3 Ativação mensal/qualificação", 2)
add_para(doc, "A configuração vigente exige R$ 79,00 em compras pessoais por mês e prevê dez dias de tolerância após o fechamento do ciclo.")
add_callout(doc, "Comportamento atual do sistema", "Essa regra está salva e aparece nas configurações, porém não é consultada pela rotina que gera comissões. Hoje, o critério efetivo para receber uma nova comissão é o status cadastral “ativo”; a qualificação mensal ainda não bloqueia, retém ou cancela comissões automaticamente.", "danger")
add_para(doc, "Até que a automação seja implementada, trate a ativação mensal como regra comercial/operacional que precisa de controle manual — e não afirme ao embaixador que o sistema já faz esse bloqueio sozinho.")

doc.add_page_break()

add_heading(doc, "6. Ciclo de vida da comissão", 1)
add_table(doc, ["Status da comissão", "Quando acontece", "Pode pagar?"], [
    ("Aguardando entrega", "Pedido ainda não chegou a entregue/finalizado.", "Não"),
    ("Aguardando pagamento", "Pedido entregue/finalizado, mas conferência financeira ainda não confirmada.", "Não"),
    ("Liberada", "Entrega/finalização e pagamento do pedido confirmados.", "Sim"),
    ("Paga", "Foi incluída em um repasse Pix concluído.", "Já paga"),
    ("Cancelada", "Pedido cancelado.", "Não"),
    ("Estornada", "Comissão revertida por procedimento financeiro.", "Não"),
], [2200, 5300, 1740])
add_bullets(doc, [
    "O pedido cancelado cancela as comissões relacionadas que ainda não foram pagas.",
    "A liberação depende de dois fatos: situação logística concluída e conferência financeira confirmada.",
    "Dados financeiros já gerados são imutáveis; correções exigem fluxo administrativo auditado.",
])

add_heading(doc, "7. Pagamentos e saques", 1)
add_callout(doc, "Regra vigente", "Repasse mensal por Pix, a partir de R$ 50,00 em comissões liberadas.", "success")
add_heading(doc, "Como o pagamento é feito", 2)
add_steps(doc, [
    ("Listar comissões liberadas. ", "O financeiro vê o saldo agrupado por embaixador e a chave Pix mascarada."),
    ("Selecionar itens. ", "É possível escolher todas ou apenas algumas comissões liberadas do mesmo embaixador."),
    ("Validar o mínimo. ", "A soma selecionada precisa atingir R$ 50,00."),
    ("Registrar referência. ", "O administrador informa um identificador único do Pix e, se quiser, observações."),
    ("Concluir. ", "O pagamento e seus itens são gravados; as comissões passam imediatamente para “paga”."),
])
add_bullets(doc, [
    "Pagamento abaixo do mínimo é possível apenas com override administrativo e justificativa auditada.",
    "A referência do pagamento é única, impedindo repetição acidental do mesmo repasse.",
    "O sistema registra snapshots do nome, CPF mascarado e chave Pix mascarada usados no pagamento.",
    "O embaixador acompanha o histórico em “Meus Pagamentos & Saques” e pode acessar comprovantes quando anexados.",
])
add_callout(doc, "Atenção operacional", "A frequência “mensal” define o ciclo de conferência, mas o pagamento é iniciado manualmente pelo administrador; não há repasse automático agendado na implementação atual.", "warning")

add_heading(doc, "Checklist financeiro", 2)
add_bullets(doc, [
    "Confirmar que todas as comissões selecionadas estão liberadas.",
    "Conferir titularidade e validade da chave Pix antes do repasse.",
    "Usar uma referência rastreável (ex.: PIX-2026-000123).",
    "Anexar comprovante e observação quando aplicável.",
    "Usar override abaixo do mínimo somente como exceção justificada.",
])

doc.add_page_break()

add_heading(doc, "8. Cadastro administrativo: passo a passo", 1)
add_steps(doc, [
    ("Abrir Novo Embaixador. ", "Apenas administradores ativos podem concluir o cadastro."),
    ("Preencher identificação. ", "Nome completo, nome de exibição, CPF válido e e-mail de contato."),
    ("Preencher contato e endereço. ", "Telefone com DDD é obrigatório para gerar a senha inicial."),
    ("Cadastrar Pix. ", "Escolher o tipo de chave e informar chave no nome da pessoa cadastrada."),
    ("Escolher o plano. ", "Usar o plano padrão vigente, salvo decisão operacional aprovada."),
    ("Vincular patrocinador. ", "Selecionar “Indicado por” quando o embaixador fizer parte de uma rede."),
    ("Definir status inicial. ", "Pendente impede login; ativo permite login imediato após a conta ser criada."),
    ("Salvar e entregar credenciais. ", "Enviar código, link e orientação para trocar a senha no primeiro acesso."),
])
add_heading(doc, "Dados e restrições importantes", 2)
add_bullets(doc, [
    "CPF e e-mail de contato não podem duplicar outro embaixador.",
    "Código, username e número são gerados automaticamente e não podem ser alterados depois.",
    "O código segue bryzaNN, em letras minúsculas.",
    "Não é permitido apagar fisicamente um embaixador; use inativo ou bloqueado.",
    "A rede impede auto-patrocínio e ciclos.",
    "Alterar o plano afeta apenas operações futuras; o histórico permanece congelado.",
])

add_heading(doc, "Primeiro acesso", 2)
add_steps(doc, [
    "Entrar com código Bryza, e-mail, CPF ou telefone.",
    "Usar como senha inicial o telefone com DDD, somente números.",
    "Informar o CPF para validação.",
    "Criar e confirmar uma senha pessoal de pelo menos oito caracteres.",
    "Acessar o dashboard do embaixador.",
])
add_callout(doc, "Suporte", "Se o embaixador já possui usuário e perdeu o acesso, o administrador pode redefinir a senha para o telefone cadastrado e forçar uma nova troca no primeiro acesso.", "info")

add_heading(doc, "9. O que existe no painel do embaixador", 1)
add_table(doc, ["Área", "Para que serve"], [
    ("Dashboard", "Resumo de vendas, comissões aguardando/disponíveis/pagas, bônus e clientes indicados."),
    ("Meu link", "Código, link completo, botão de copiar, compartilhamento por WhatsApp e QR Code."),
    ("Indicações", "Clientes indicados, origem, status e sinais de ativação por compra confirmada."),
    ("Vendas", "Pedidos atribuídos ao embaixador e seus valores/status."),
    ("Comissões", "Extrato por pedido, percentual ou bônus, valor e status."),
    ("Minha rede", "Embaixadores descendentes até o terceiro nível, com filtros por nível."),
    ("Pagamentos", "Histórico de repasses e acesso a comprovantes disponíveis."),
    ("Perfil", "Dados pessoais, endereço, localização, foto e informações Pix permitidas."),
    ("Materiais", "Área destinada a guias e artes; downloads ainda aparecem como “em breve”."),
    ("Calculadora", "Simulação de ganhos segundo o plano 4–2–1."),
], [2000, 7240])

doc.add_page_break()

add_heading(doc, "10. Como ensinar o programa", 1)
add_heading(doc, "Roteiro de apresentação em 5 minutos", 2)
add_steps(doc, [
    ("Comece pelo propósito. ", "“A Bryza recompensa quem apresenta clientes e ajuda a ampliar uma rede de recomendação.”"),
    ("Mostre o link. ", "Abra a área Meu Link, copie o endereço e mostre o QR Code."),
    ("Explique 4–2–1 com R$ 100. ", "R$ 4 direto, R$ 2 no segundo nível e R$ 1 no terceiro."),
    ("Explique a validação. ", "A comissão só libera depois da entrega e da confirmação do pagamento."),
    ("Explique o repasse. ", "A partir de R$ 50 liberados, o financeiro paga por Pix no ciclo mensal."),
    ("Feche com os cuidados. ", "Sem autoindicação, dados corretos e divulgação verdadeira."),
])

add_heading(doc, "Perguntas de checagem", 2)
add_bullets(doc, [
    "A pessoa consegue dizer quem é nível 1, 2 e 3 em um exemplo real?",
    "Ela entendeu que comissão criada não é o mesmo que comissão liberada?",
    "Ela sabe onde copiar o link e acompanhar vendas e comissões?",
    "Ela entendeu o mínimo de R$ 50 e o ciclo mensal?",
    "Ela sabe que o código recebido após uma compra não significa acesso automático ao painel?",
])

add_heading(doc, "Script de onboarding", 2)
add_callout(doc, "Mensagem sugerida", "“Bem-vindo ao Programa de Embaixadores Bryza. Seu código é [CÓDIGO] e seu link é [LINK]. Use seu telefone com DDD, apenas números, como senha inicial; no primeiro acesso você criará uma senha pessoal. Compartilhe o link para que as vendas sejam atribuídas corretamente. Você poderá acompanhar indicações, vendas, comissões e pagamentos no painel.”", "info")

add_heading(doc, "11. Boas práticas e conduta", 1)
add_bullets(doc, [
    "Compartilhar sempre o link oficial completo; evitar digitar códigos de memória.",
    "Não prometer renda garantida, valor fixo de ganhos ou prazo de pagamento fora das regras vigentes.",
    "Não fazer autoindicação nem usar dados de terceiros para simular venda ou ativação.",
    "Confirmar que o cliente sabe que a compra está vinculada a uma indicação.",
    "Não expor CPF, chave Pix, telefone ou dados de clientes em grupos e redes sociais.",
    "Usar materiais e mensagens oficiais da Bryza.",
    "Manter telefone, endereço e Pix atualizados no perfil.",
    "Encaminhar divergências ao suporte com código do embaixador e número do pedido, sem dados sensíveis desnecessários.",
])

add_heading(doc, "12. Diagnóstico rápido", 1)
add_table(doc, ["Sintoma", "Causa provável", "O que verificar"], [
    ("Link não abre", "Código inexistente/inativo ou configuração de ambiente", "Status do embaixador, domínio e formato bryzaNN"),
    ("Compra diz indicação expirada", "Cookie/visita ausente ou fora da janela", "Abrir novamente o link original e refazer o checkout"),
    ("Venda não aparece", "Atribuição ficou com o primeiro indicador bloqueado", "Cliente e histórico de atribuição"),
    ("Comissão não foi criada", "Autoindicação, beneficiário inativo ou plano indisponível", "CPF/telefone, cadeia, status e snapshots"),
    ("Comissão não liberou", "Entrega ou conferência financeira pendente", "Status do pedido e payment check"),
    ("Não dá para pagar", "Saldo selecionado abaixo de R$ 50", "Somar liberadas ou justificar override"),
    ("Código existe, mas não entra", "Candidato automático sem usuário de autenticação", "Provisionar acesso antes de ativar"),
    ("Ativação mensal não bloqueou ganho", "Regra ainda não integrada ao cálculo", "Aplicar controle manual até a automação"),
], [2350, 3250, 3640])

doc.add_page_break()

add_heading(doc, "13. Perguntas frequentes", 1)
faq = [
    ("Quanto eu ganho em uma venda direta?", "No plano atual, 4% sobre a base comissionável do pedido."),
    ("Posso ganhar nos indicados dos meus indicados?", "Sim. O plano atual paga 2% no nível 2 e 1% no nível 3."),
    ("Quando a comissão fica disponível?", "Quando o pedido está entregue ou finalizado e o pagamento foi confirmado."),
    ("Quando recebo?", "No ciclo mensal, quando houver pelo menos R$ 50 em comissões liberadas; a operação registra o Pix manualmente."),
    ("Se eu mudar de plano, muda o passado?", "Não. Pedidos e comissões mantêm os percentuais congelados quando foram criados."),
    ("Se um nível estiver inativo, outra pessoa recebe no lugar?", "Não. O sistema pula o pagamento daquele beneficiário, sem comprimir ou redistribuir."),
    ("Posso comprar pelo meu próprio link?", "Não para gerar comissão. Autoindicação por vínculo, CPF ou telefone é bloqueada."),
    ("O bônus de R$ 7 vale em qualquer compra?", "Não. Vale uma vez, na primeira compra elegível do novo indicado, a partir de R$ 79, com entrega e pagamento confirmados e demais vínculos válidos."),
    ("A compra me transforma em embaixador?", "Ela cria um candidato e um código próprio. O acesso ao painel ainda precisa ser provisionado e ativado administrativamente."),
    ("Preciso movimentar R$ 79 todo mês?", "Essa é a regra comercial configurada, baseada em compras pessoais e com dez dias de tolerância. Hoje, porém, o sistema ainda não a aplica automaticamente ao cálculo de comissões."),
    ("Posso editar minha chave Pix?", "A configuração atual permite edição e não exige aprovação prévia, mas a chave deve pertencer à pessoa cadastrada."),
    ("Quem pode ver meus dados?", "O embaixador vê o próprio perfil. Operações administrativas sensíveis são restritas e auditadas; dados financeiros aparecem mascarados nos fluxos de pagamento."),
]
for question, answer in faq:
    p = add_para(doc, after=2)
    r = p.add_run(question)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_para(doc, answer, after=7)

add_heading(doc, "14. Checklist de lançamento/rotina", 1)
add_heading(doc, "Antes de apresentar o programa", 2)
add_bullets(doc, [
    "Confirmar se o programa continua ativo e se o plano padrão é 4–2–1.",
    "Confirmar valores de ativação, bônus, mínimo e frequência de pagamento.",
    "Validar o domínio do link e o funcionamento do QR Code.",
    "Alinhar o procedimento de provisionamento para candidatos automáticos.",
    "Definir o responsável pelo controle manual da ativação mensal.",
])
add_heading(doc, "Toda semana", 2)
add_bullets(doc, [
    "Revisar pedidos entregues com conferência financeira pendente.",
    "Investigar comissões presas ou divergentes.",
    "Acompanhar candidatos pendentes que precisam de ativação de acesso.",
    "Verificar alterações de Pix e dados cadastrais.",
])
add_heading(doc, "No fechamento mensal", 2)
add_bullets(doc, [
    "Apurar qualificação mensal conforme a regra comercial vigente.",
    "Listar comissões liberadas por embaixador.",
    "Conferir o mínimo de R$ 50 e exceções aprovadas.",
    "Executar Pix, registrar referência e anexar comprovantes.",
    "Reconciliar pagamentos, comissões pagas e trilha de auditoria.",
])

doc.add_page_break()

add_heading(doc, "15. Glossário", 1)
add_table(doc, ["Termo", "Definição operacional"], [
    ("Embaixador atribuído", "Pessoa cuja indicação está vinculada ao cliente/pedido e ocupa o nível 1."),
    ("Patrocinador", "Embaixador acima de outro na árvore; pode receber em níveis superiores."),
    ("Candidato", "Cadastro automático criado pela compra, ainda pendente de acesso."),
    ("Atribuição", "Vínculo entre cliente e embaixador indicador."),
    ("Janela", "Período máximo entre a visita rastreada e o uso válido da indicação."),
    ("Snapshot", "Cópia imutável da regra e dos valores vigentes no momento do pedido."),
    ("Base comissionável", "Valor sobre o qual o percentual é aplicado."),
    ("Comissão liberada", "Valor validado por entrega/finalização e confirmação financeira."),
    ("Override", "Exceção administrativa auditada, como pagar abaixo do mínimo."),
    ("Autoindicação", "Tentativa de receber comissão sobre a própria compra; é bloqueada."),
], [2100, 7140])

add_heading(doc, "16. Pontos de governança e evolução", 1)
add_callout(doc, "Prioridade alta", "Criar um fluxo administrativo único para transformar candidatos automáticos em usuários ativos, com criação da conta, senha inicial e confirmação de dados, sem depender de recadastro duplicado.", "danger")
add_callout(doc, "Prioridade alta", "Integrar a qualificação mensal de R$ 79 ao motor de comissões, definindo explicitamente o efeito: bloquear geração, manter pendente, liberar após regularização ou perder a comissão. Hoje esse efeito não existe no código.", "danger")
add_bullets(doc, [
    "Definir uma política formal para troca de indicador e quais estados permitem reatribuição.",
    "Definir SLA de conferência financeira, pagamento e envio de comprovante.",
    "Publicar materiais oficiais na área Materiais.",
    "Transformar a frequência mensal em calendário operacional visível.",
    "Revisar periodicamente se cookie, janela configurada e mensagens públicas permanecem sincronizados.",
])

add_heading(doc, "Base de verificação", 2)
add_para(doc, "Este playbook foi elaborado a partir das configurações vigentes consultadas no ambiente Bryza em 21/07/2026 e da implementação local do módulo de embaixadores: migrations de banco de dados, rotinas de atribuição/comissão/pagamento e telas administrativas e do portal do embaixador.")
add_para(doc, "Sempre que o plano ou as configurações forem alterados, atualize primeiro o quadro “Resumo executivo”, os exemplos financeiros, o roteiro de explicação e as perguntas frequentes.")

# Final metadata and save
doc.core_properties.title = "Playbook do Programa de Embaixadores Bryza"
doc.core_properties.subject = "Operação, comissões, ativação, rede, pagamentos e comunicação"
doc.core_properties.author = "Bryza"
doc.core_properties.keywords = "Bryza, embaixadores, comissões, indicação, multinível, playbook"
doc.save(OUT)
print(str(OUT))
